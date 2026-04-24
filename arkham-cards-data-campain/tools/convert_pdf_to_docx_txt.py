import argparse
import sys
from pathlib import Path

import fitz  # PyMuPDF


def extract_page_blocks(page):
    blocks = []
    for block in page.get_text("blocks"):
        x0, y0, x1, y1, text, *_ = block
        cleaned = text.strip()
        if cleaned:
            blocks.append((y0, x0, cleaned))
    blocks.sort(key=lambda item: (round(item[0], 1), round(item[1], 1)))
    return [text for _, _, text in blocks]


def extract_pdf_pages(pdf_path):
    doc = fitz.open(pdf_path)
    pages = []
    for page in doc:
        pages.append(extract_page_blocks(page))
    return pages


def build_txt_content(pages, include_page_markers=True):
    chunks = []
    for index, blocks in enumerate(pages, start=1):
        page_text = "\n\n".join(blocks).strip()
        if include_page_markers:
            chunks.append(f"=== PAGE {index} ===\n{page_text}")
        else:
            chunks.append(page_text)
    return "\n\n".join(chunks).strip() + "\n"


def write_docx(pages, output_docx):
    try:
        from docx import Document
    except ModuleNotFoundError as error:
        raise RuntimeError(
            "Missing dependency 'python-docx'. Install with: pip install python-docx"
        ) from error

    document = Document()
    for page_index, blocks in enumerate(pages, start=1):
        document.add_heading(f"Page {page_index}", level=2)
        for block in blocks:
            for line in block.splitlines():
                text = line.strip()
                if text:
                    document.add_paragraph(text)
            document.add_paragraph("")
        if page_index < len(pages):
            document.add_page_break()
    document.save(output_docx)


def convert_pdf(pdf_path, output_dir, include_page_markers=True, quiet=False):
    pages = extract_pdf_pages(pdf_path)
    output_dir.mkdir(parents=True, exist_ok=True)

    output_txt = output_dir / f"{pdf_path.stem}.txt"
    output_docx = output_dir / f"{pdf_path.stem}.docx"

    txt_content = build_txt_content(pages, include_page_markers=include_page_markers)
    output_txt.write_text(txt_content, encoding="utf-8")
    write_docx(pages, output_docx)

    if not quiet:
        print(f"[OK] {pdf_path}")
        print(f"  TXT : {output_txt}")
        print(f"  DOCX: {output_docx}")


def find_pdf_files(explicit_files, input_dir):
    if explicit_files:
        return [Path(path) for path in explicit_files]
    return sorted(Path(input_dir).glob("*.pdf"))


def parse_args():
    parser = argparse.ArgumentParser(
        description="Convert PDF files to TXT and DOCX."
    )
    parser.add_argument(
        "pdf_files",
        nargs="*",
        help="Optional list of PDF files. If omitted, all PDFs in --input-dir are used.",
    )
    parser.add_argument(
        "--input-dir",
        default=".",
        help="Directory scanned for PDFs when no file is provided.",
    )
    parser.add_argument(
        "--output-dir",
        default=None,
        help="Output directory. Defaults to each PDF's own directory.",
    )
    parser.add_argument(
        "--no-page-markers",
        action="store_true",
        help="Do not add '=== PAGE n ===' markers to TXT output.",
    )
    parser.add_argument(
        "--quiet",
        action="store_true",
        help="Suppress per-file success logs.",
    )
    return parser.parse_args()


def main():
    args = parse_args()
    pdf_files = find_pdf_files(args.pdf_files, args.input_dir)

    if not pdf_files:
        print("No PDF files found.")
        return 1

    failures = 0
    for pdf_file in pdf_files:
        try:
            if not pdf_file.exists():
                raise FileNotFoundError(f"File not found: {pdf_file}")
            if pdf_file.suffix.lower() != ".pdf":
                raise ValueError(f"Not a PDF file: {pdf_file}")

            output_dir = Path(args.output_dir) if args.output_dir else pdf_file.parent
            convert_pdf(
                pdf_file,
                output_dir,
                include_page_markers=not args.no_page_markers,
                quiet=args.quiet,
            )
        except Exception as error:
            failures += 1
            print(f"[ERROR] {pdf_file}: {error}")

    if failures:
        print(f"Completed with {failures} failure(s).")
        return 1

    print(f"Completed successfully. Converted {len(pdf_files)} file(s).")
    return 0


if __name__ == "__main__":
    sys.exit(main())
